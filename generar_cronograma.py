import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores como strings HEX ──────────────────────────────────────────────
C_AZUL_OSC   = "00497A"
C_AZUL_MED   = "0070B0"
C_AMARILLO   = "FFB300"
C_GRIS_FOND  = "F0F4F8"
C_GRIS_CLAR  = "D0D8E0"
C_VERDE      = "D4EDDA"
C_VERDE_TXT  = "156E3A"
C_AMBAR_BG   = "FFF0B8"
C_AMBAR_TXT  = "925E00"
C_BLANCO     = "FFFFFF"
C_NEGRO      = "1A1A2E"
C_TEXTO      = "333344"
C_MUTED      = "666688"
C_NOTA       = "666688"

ZONA_HEX   = ["3B82F6","8B5CF6","10B981","F59E0B","EF4444"]
ZONA_RGB   = [
    RGBColor(0x3B,0x82,0xF6),
    RGBColor(0x8B,0x5C,0xF6),
    RGBColor(0x10,0xB9,0x81),
    RGBColor(0xF5,0x9E,0x0B),
    RGBColor(0xEF,0x44,0x44),
]

# ── Datos ─────────────────────────────────────────────────────────────────
project_data = [
    {
        "zone": "Zona 1: Cochera", "emoji": "Z1",
        "tasks": ["Trazo/Excava","Zapata/Dado","Ancla/Placa","Columna CM","Viga V-1","Pérgola",
                  "Terracerías","Firme","Aplanados","Instalaciones","Pintura base","Piso",
                  "Herrería","Detallado","Limpieza","Entrega"],
        "details": [
            ("Trazo y excavación",   "Trazo topográfico y excavación manual para cepas de 1.00×1.00m según plano E-01."),
            ("Zapatas",              "Vaciado de plantilla de concreto pobre. Armado de acero y colado de zapatas aisladas D-1/D-2."),
            ("Anclajes",             "Presentación, nivelación y fijación de anclas de 5/8\" y placas base PL-1 y PL-2."),
            ("Columnas",             "Izaje, montaje y plomeo de columnas metálicas de tubo cuadrado (CM, 14.88 kg/m) con soldadura a penetración."),
            ("Vigas principales",    "Montaje y soldadura de vigas principales V-1 (I Rectangular IPS 5\")."),
            ("Pérgola",              "Colocación de vigas secundarias V-2, V-4 y estándar V-3 para el entramado del pergolado metálico."),
            ("Terracerías",          "Excavación menor, relleno compactado con material de banco y nivelación de terreno para estacionamiento."),
            ("Firme",                "Colocación de malla electrosoldada y vaciado de concreto para firme de 10 cm de espesor."),
            ("Aplanados",            "Resane y aplicación de aplanados base cemento-arena a plomo y regla en muros colindantes."),
            ("Instalaciones",        "Ranurado y tendido de poliducto para salidas eléctricas de iluminación en cubierta de pérgola."),
            ("Pintura base",         "Aplicación de sellador acrílico en muros y primario anticorrosivo en estructura metálica."),
            ("Piso",                 "Instalación de acabado en piso (cerámico para tráfico pesado o concreto estampado)."),
            ("Herrería",             "Instalación de portón de herrería corredizo o abatible y detallado de fachada principal."),
            ("Detallado",            "Aplicación de pintura tipo esmalte en estructura metálica y pintura vinílica en muros exteriores."),
            ("Limpieza",             "Retiro de cimbra restante, recolección de escombro y limpieza general del área de estacionamiento."),
            ("Entrega",              "Recorrido de revisión de soldaduras, acabados, pintura e iluminación con el cliente."),
        ],
    },
    {
        "zone": "Zona 2: Baño Principal", "emoji": "Z2",
        "tasks": ["Demolición","Muros","Castillos","Inst. Hidro","Inst. Elec","Firme",
                  "Enjarres","Yeso","Azulejo","Carpintería","Muebles","Canceles",
                  "Pintura","Limpieza","Detalles","Entrega"],
        "details": [
            ("Demolición",   "Demolición de recubrimientos cerámicos existentes, desmantelamiento de muebles de baño y clósets viejos."),
            ("Muros",        "Trazo y desplante de nuevos muros divisorios de block ligero asentado con mortero para nueva distribución."),
            ("Castillos",    "Armado, cimbrado y colado de castillos de confinamiento (K-1, K-2) para asegurar los muros nuevos."),
            ("Inst. Hidro",  "Ranurado de muros y piso para tendido de tubería de PVC sanitario y CPVC para agua fría/caliente."),
            ("Inst. Elec",   "Canalización de ductos eléctricos para contactos de lavabo, iluminación general y extractores."),
            ("Firme",        "Vaciado de firme de compresión y formación de charola impermeabilizada con pendiente en regadera."),
            ("Enjarres",     "Aplicación de enjarre rústico y repellado plomeado en muros nuevos y áreas reparadas."),
            ("Yeso",         "Aplicación de pasta fina o yeso en muros fuera de las áreas húmedas y en plafones."),
            ("Azulejo",      "Instalación de piso cerámico antiderrapante y azulejo en muros de la zona de regadera."),
            ("Carpintería",  "Ensamblaje e instalación de módulos de madera para vestidor y marcos de puertas de intercomunicación."),
            ("Muebles",      "Instalación de inodoro, lavabo sobre cubierta, mezcladoras y accesorios de grifería empotrada."),
            ("Canceles",     "Colocación de cancel de baño de cristal templado, espejos y placas eléctricas/luminarias."),
            ("Pintura",      "Aplicación de pintura vinílica lavable especial para zonas húmedas en vestidor y techo."),
            ("Limpieza",     "Limpieza profunda con ácido muriático diluido para retiro de manchas de boquilla y polvo fino."),
            ("Detalles",     "Aplicación de sellador de silicón anti-hongos en juntas y revisión de presión en instalaciones hidro."),
            ("Entrega",      "Pruebas de funcionamiento eléctrico y de flujo de agua. Entrega de la zona habitable."),
        ],
    },
    {
        "zone": "Zona 3: Recámara 2", "emoji": "Z3",
        "tasks": ["Excavación","Cimentación","Muros PB","Castillos PB","Cimbra PB","Losa PB",
                  "Muros PA","Anclaje PA","Cimbra PA","Acero PA","Losa PA","Ranurados",
                  "Aplanados","Firmes","Carpintería","Entrega"],
        "details": [
            ("Excavación",    "Trazo y excavación de zanjas para cimentación corrida (ZC-1) con profundidad según mecánica de suelos."),
            ("Cimentación",   "Vaciado de plantilla, habilitado de acero, cimbrado y colado de dalas de desplante (DL-1, DL-2)."),
            ("Muros PB",      "Desplante y levantamiento de muros de carga de block/tabique rojo recocido en PB a plomo y nivel."),
            ("Castillos PB",  "Armado, cimbrado y colado de castillos (K-1) para confinar los muros perimetrales e interiores de PB."),
            ("Cimbra PB",     "Apuntalamiento y cimbrado de madera alomada para recibir la losa nervada de entrepiso."),
            ("Losa PB",       "Colocación de casetones de poliestireno, armado de trabes (T-1 a T-3) y colado de capa de compresión."),
            ("Muros PA",      "Periodo de curado de losa y desplante de nuevos muros de mampostería en la nueva Planta Alta."),
            ("Anclaje PA",    "Ejecución de detalle estructural de anclaje dentado para conectar castillos nuevos con estructura existente."),
            ("Cimbra PA",     "Apuntalamiento, nivelación y colocación de tarimas/cimbra para la losa de azotea en Planta Alta."),
            ("Acero PA",      "Habilitado de acero para trabes perimetrales (T-6 a T-9), nervaduras (N-7 a N-11) y casetones."),
            ("Losa PA",       "Colado de losa nervada de azotea con concreto premezclado, pretiles y formación de pendientes pluviales."),
            ("Ranurados",     "Ranurado de muros para bajantes pluviales, tubería sanitaria, hidráulica y ductos eléctricos en ambos niveles."),
            ("Aplanados",     "Aplicación de enjarre a plomo y regla en muros interiores y aplanado impermeable en fachadas."),
            ("Firmes",        "Vaciado de firmes niveladores en PB y PA. Preparación de superficie para pisos de porcelanato."),
            ("Carpintería",   "Instalación de piso, zoclos, puertas de intercomunicación de tambor/madera y frentes de clóset."),
            ("Entrega",       "Aplicación de pasta fina, pintura vinílica, cableado final, prueba de luminarias y limpieza de entrega."),
        ],
    },
    {
        "zone": "Zona 4: Recámara 1", "emoji": "Z4",
        "tasks": ["Demolición","Muros","Castillos","Inst. Hidro","Inst. Elec","Firmes",
                  "Enjarres","Yeso","Piso","Azulejo","Carpintería","Muebles",
                  "Pintura","Cancelería","Limpieza","Entrega"],
        "details": [
            ("Demolición",   "Desmontaje de puertas, retiro de piso cerámico antiguo y demolición parcial de muros para ampliación de vanos."),
            ("Muros",        "Levantamiento de muros divisorios de block para integrar el baño completo a la habitación."),
            ("Castillos",    "Armado y colado de dalas de cerramiento y castillos para refuerzo de nuevos vanos de puertas y ventanas."),
            ("Inst. Hidro",  "Ruptura parcial de firme existente para conectar nueva red de drenaje sanitario y alimentación de agua."),
            ("Inst. Elec",   "Tendido de poliducto flexible corrugado en muros y losa para nuevos apagadores y contactos."),
            ("Firmes",       "Relleno de zanjas, compactación y vaciado de firme autonivelante para recibir piso nuevo."),
            ("Enjarres",     "Aplicación de repellado base en áreas de muros intervenidos para emparejar con plomos existentes."),
            ("Yeso",         "Aplicación de acabado de yeso en muros interiores de la recámara y plafones."),
            ("Piso",         "Instalación de piso cerámico o formato tipo madera en recámara, respetando juntas de dilatación."),
            ("Azulejo",      "Colocación de recubrimiento cerámico en formato grande en muros de zona de regadera y lavabo."),
            ("Carpintería",  "Instalación de clóset a medida y montaje de puerta principal de la recámara con chambranas."),
            ("Muebles",      "Instalación de muebles sanitarios, monomandos, regadera de teléfono y accesorios de baño."),
            ("Pintura",      "Aplicación de sellador y dos manos de pintura vinílica acrílica en muros y plafones."),
            ("Cancelería",   "Instalación de ventanas de aluminio con mosquiteros y cancel corredizo en cristal inastillable."),
            ("Limpieza",     "Limpieza general de obra, retiro de plásticos protectores y cintas de enmascarar."),
            ("Entrega",      "Verificación de acabados, revisión de caída de agua en regadera y entrega funcional."),
        ],
    },
    {
        "zone": "Zona 5: Sala TV", "emoji": "Z5",
        "tasks": ["Excavación","Cimentación","Muros PB","Castillos PB","Cimbra PB","Losa PB",
                  "Muros PA","Anclaje PA","Cimbra PA","Acero PA","Losa PA","Ranurados",
                  "Aplanados","Firmes","Yeso","Entrega"],
        "details": [
            ("Excavación",   "Excavación de cepas para extensión de cimentación corrida y zapatas según plano estructural E-01."),
            ("Cimentación",  "Armado de acero estructural, cimbrado y colado de cimiento de mampostería/concreto y dalas perimetrales."),
            ("Muros PB",     "Desplante de muros de carga de ladrillo o block pesado en el área de expansión de la Sala TV en PB."),
            ("Castillos PB", "Colado de castillos (K-1) para dar rigidez a la estructura nueva e integración con muros existentes."),
            ("Cimbra PB",    "Montaje de andamios y cimbra de madera o metálica para soportar el colado de la losa de entrepiso."),
            ("Losa PB",      "Armado de sistema de losa nervada, colocación de casetones aligerantes y vaciado de concreto premezclado."),
            ("Muros PA",     "Desplante de muros perimetrales en Planta Alta para el área de cuarto de juegos."),
            ("Anclaje PA",   "Ejecución técnica de anclaje dentado de mampostería para integrar muros nuevos con la casa actual."),
            ("Cimbra PA",    "Preparación estructural de la cubierta final (azotea) mediante apuntalamiento perimetral y central."),
            ("Acero PA",     "Armado y amarre de acero para trabes de carga principales (T-6 a T-9) y nervaduras estructurales secundarias."),
            ("Losa PA",      "Vaciado de concreto en losa de azotea, construcción de pretiles de protección y pendientes para lluvia."),
            ("Ranurados",    "Ranurado con esmeril para ocultar preparaciones de cableado para TV, internet, contactos y audio."),
            ("Aplanados",    "Aplicación de mortero cemento-arena en fachadas y muros interiores para recibir acabado."),
            ("Firmes",       "Vaciado de firmes de concreto en PB y PA. Aplicación de impermeabilizante en azotea nueva."),
            ("Yeso",         "Aplicación de pasta fina en interiores e instalación de piso porcelánico en gran formato."),
            ("Entrega",      "Instalación de luminarias LED, pintura vinílica, retiro de andamios exteriores y limpieza final."),
        ],
    },
]

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper().replace('#',''))
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=8, rgb=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if rgb:
        run.font.color.rgb = rgb

def add_border(table, color="AAAAAA"):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),  'single')
        b.set(qn('w:sz'),   '4')
        b.set(qn('w:space'),'0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tbl.tblPr.append(tblBorders)

# ── Documento ─────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width    = Cm(29.7)
section.page_height   = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.orientation   = 1  # landscape

# ── Portada ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CRONOGRAMA MAESTRO — AMPLIACIÓN SAUCES")
r.bold = True; r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x00,0x49,0x7A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Especificaciones Técnicas por Semana  ·  5 Zonas  ·  16 Semanas")
r2.font.size = Pt(11)
r2.font.color.rgb = RGBColor(0x00,0x70,0xB0)

doc.add_paragraph()

# Leyenda
lp = doc.add_paragraph()
lp.add_run("Leyenda:   ").bold = True
for label, rgb in [
    ("Completado (Sem 1-4)", RGBColor(0x15,0x6E,0x3A)),
    ("   Semana actual (Sem 5)", RGBColor(0x92,0x5E,0x00)),
    ("   Pendiente", RGBColor(0x44,0x44,0x55)),
]:
    rr = lp.add_run(f"■ {label}")
    rr.font.color.rgb = rgb
    rr.font.size = Pt(9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — GANTT GENERAL
# ══════════════════════════════════════════════════════════════════════════
h1 = doc.add_heading("1.  Tabla Gantt General", level=1)
h1.runs[0].font.color.rgb = RGBColor(0x00,0x49,0x7A)

gantt = doc.add_table(rows=1 + len(project_data), cols=1 + 16)
gantt.alignment = WD_TABLE_ALIGNMENT.CENTER
add_border(gantt)

# Encabezado
hrow = gantt.rows[0]
set_cell_bg(hrow.cells[0], C_AZUL_OSC)
cell_text(hrow.cells[0], "ZONA / SEMANA", bold=True, size=8,
          rgb=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.LEFT)
for i in range(1, 17):
    set_cell_bg(hrow.cells[i], C_AZUL_OSC)
    cell_text(hrow.cells[i], f"S{i:02d}", bold=True, size=7, rgb=RGBColor(0xFF,0xFF,0xFF))

# Filas
for zi, rd in enumerate(project_data):
    row = gantt.rows[zi + 1]
    zrgb = ZONA_RGB[zi]
    set_cell_bg(row.cells[0], C_GRIS_FOND)
    cell_text(row.cells[0], f"{rd['emoji']}  {rd['zone']}",
              bold=True, size=8, rgb=zrgb, align=WD_ALIGN_PARAGRAPH.LEFT)
    for wi in range(16):
        cell = row.cells[wi + 1]
        task_name = rd["tasks"][wi]
        if wi < 4:
            set_cell_bg(cell, C_VERDE)
            cell_text(cell, task_name, size=6.5, rgb=RGBColor(0x15,0x6E,0x3A))
        elif wi == 4:
            set_cell_bg(cell, C_AMBAR_BG)
            cell_text(cell, task_name, bold=True, size=6.5, rgb=RGBColor(0x92,0x5E,0x00))
        else:
            set_cell_bg(cell, C_GRIS_CLAR if wi % 2 == 0 else C_BLANCO)
            cell_text(cell, task_name, size=6.5, rgb=RGBColor(0x44,0x44,0x55))

# Anchos
for row in gantt.rows:
    row.cells[0].width = Cm(3.8)
    for i in range(1, 17):
        row.cells[i].width = Cm(1.45)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — DETALLE POR ZONA
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h2 = doc.add_heading("2.  Especificaciones Técnicas por Zona", level=1)
h2.runs[0].font.color.rgb = RGBColor(0x00,0x49,0x7A)

for zi, rd in enumerate(project_data):
    zrgb  = ZONA_RGB[zi]
    zhex  = ZONA_HEX[zi]

    zh = doc.add_heading(f"  {rd['zone']}", level=2)
    zh.runs[0].font.color.rgb = zrgb

    tbl = doc.add_table(rows=1 + 16, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_border(tbl)

    # Encabezado
    hrw = tbl.rows[0]
    for ci, lbl in enumerate(["Semana", "Actividad", "Descripción Técnica"]):
        set_cell_bg(hrw.cells[ci], C_AZUL_OSC)
        cell_text(hrw.cells[ci], lbl, bold=True, size=9, rgb=RGBColor(0xFF,0xFF,0xFF))

    # Filas
    for wi, (ttl, dsc) in enumerate(rd["details"]):
        r = tbl.rows[wi + 1]
        bg = C_GRIS_FOND if wi % 2 == 0 else C_BLANCO
        set_cell_bg(r.cells[0], bg)
        set_cell_bg(r.cells[1], bg)
        set_cell_bg(r.cells[2], bg)
        cell_text(r.cells[0], f"Sem {wi+1:02d}", bold=True, size=9, rgb=zrgb)
        cell_text(r.cells[1], f"• {ttl}", bold=True, size=9,
                  rgb=RGBColor(0x1A,0x1A,0x2E), align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_text(r.cells[2], dsc, size=8.5,
                  rgb=RGBColor(0x33,0x33,0x44), align=WD_ALIGN_PARAGRAPH.LEFT)

    # Anchos
    for r2 in tbl.rows:
        r2.cells[0].width = Cm(1.8)
        r2.cells[1].width = Cm(4.5)
        r2.cells[2].width = Cm(20.0)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — CHECKLIST DE ENTREGA
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h3 = doc.add_heading("3.  Checklist de Verificación — Entrega por Zona", level=1)
h3.runs[0].font.color.rgb = RGBColor(0x00,0x49,0x7A)

ck = doc.add_table(rows=1 + len(project_data), cols=5)
ck.alignment = WD_TABLE_ALIGNMENT.CENTER
add_border(ck)

for ci, lbl in enumerate(["Zona","Total tareas","Tareas completadas","% Avance","Observaciones / Notas"]):
    set_cell_bg(ck.rows[0].cells[ci], C_AZUL_OSC)
    cell_text(ck.rows[0].cells[ci], lbl, bold=True, size=9, rgb=RGBColor(0xFF,0xFF,0xFF))

for zi, rd in enumerate(project_data):
    r = ck.rows[zi + 1]
    bg = C_GRIS_FOND if zi % 2 == 0 else C_BLANCO
    zrgb = ZONA_RGB[zi]
    for ci in range(5):
        set_cell_bg(r.cells[ci], bg)
    cell_text(r.cells[0], f"{rd['zone']}", bold=True, size=9,
              rgb=zrgb, align=WD_ALIGN_PARAGRAPH.LEFT)
    cell_text(r.cells[1], "16",  size=9, rgb=RGBColor(0x1A,0x1A,0x2E))
    cell_text(r.cells[2], "___", size=9, rgb=RGBColor(0x1A,0x1A,0x2E))
    cell_text(r.cells[3], "___%",size=9, rgb=RGBColor(0x1A,0x1A,0x2E))
    cell_text(r.cells[4], " ",   size=9, rgb=RGBColor(0x1A,0x1A,0x2E),
              align=WD_ALIGN_PARAGRAPH.LEFT)

for r in ck.rows:
    r.cells[0].width = Cm(4.5)
    r.cells[1].width = Cm(2.5)
    r.cells[2].width = Cm(3.5)
    r.cells[3].width = Cm(2.0)
    r.cells[4].width = Cm(13.0)

# Nota final
doc.add_paragraph()
np_ = doc.add_paragraph()
np_.alignment = WD_ALIGN_PARAGRAPH.LEFT
rn = np_.add_run(
    "Nota: Este documento es el respaldo físico imprimible del Cronograma Digital Interactivo. "
    "Para actualizar avances en tiempo real acceder a: "
    "https://manelv74-creator.github.io/cronograma-sauces/"
)
rn.font.size = Pt(8)
rn.italic = True
rn.font.color.rgb = RGBColor(0x66,0x66,0x88)

# ── Guardar ───────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Cronograma_Sauces.docx')
doc.save(out)
print(f"Documento generado: {out}")
